# src/agents/labeler.py
from parsers.bank_statement_parser import parse_bank_statement_file, rule_based_labeling
from db.mongo_client import insert_labeled_document
from utils.logger import logger
import os
import csv
from docx import Document

def export_labeled_csv(labeled_rows, out_path):
    keys = ["line_id", "date", "description", "debit", "credit", "balance", "category", "raw"]
    with open(out_path, "w", newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        for r in labeled_rows:
            row = {k: r.get(k) for k in keys}
            writer.writerow(row)

def export_labeled_docx(labeled_rows, out_path):
    doc = Document()
    doc.add_heading('Labeled Bank Statement', level=1)
    table = doc.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    hdr[0].text = 'Line'
    hdr[1].text = 'Date'
    hdr[2].text = 'Description'
    hdr[3].text = 'Debit'
    hdr[4].text = 'Credit'
    hdr[5].text = 'Balance'
    hdr[6].text = 'Category'
    for r in labeled_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.get('line_id', ''))
        row_cells[1].text = str(r.get('date', '') or '')
        row_cells[2].text = str(r.get('description', '') or '')
        row_cells[3].text = str(r.get('debit', '') or '')
        row_cells[4].text = str(r.get('credit', '') or '')
        row_cells[5].text = str(r.get('balance', '') or '')
        row_cells[6].text = str(r.get('category', '') or '')
    doc.save(out_path)

def run_labeler(payload: dict) -> dict:
    task = payload.get("task", {})
    args = task.get("args", {})
    ctx = payload.get("context", {})
    doc_path = args.get("doc_id")
    if not doc_path:
        return {"status": "error", "message": "doc_id (path) required"}
    parsed = parse_bank_statement_file(doc_path)
    labeled = parsed.get("rows", [])
    # Optionally call LLM for ambiguous rows (not implemented here; placeholder)
    # Save labeled JSON to DB and disk
    record = {
        "document_path": doc_path,
        "labels": labeled,
        "labeler_version": "v0.2"
    }
    saved = insert_labeled_document(record)
    os.makedirs("datasets/labeled_data", exist_ok=True)
    out_csv = os.path.join("datasets", "labeled_data", os.path.basename(doc_path) + ".csv")
    out_docx = os.path.join("outputs", "labeled_docs", os.path.basename(doc_path) + ".docx")
    os.makedirs(os.path.dirname(out_docx), exist_ok=True)
    export_labeled_csv(labeled, out_csv)
    export_labeled_docx(labeled, out_docx)
    logger.info(f"Labeler: saved labeled csv {out_csv} and docx {out_docx}")
    # Return labels so orchestrator can store them in shared memory
    return {"status": "ok", "labeled_count": len(labeled), "csv": out_csv, "docx": out_docx, "labels": labeled}
